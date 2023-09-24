import shutil
import telebot
import requests
from telebot import types
import os
from bs4 import BeautifulSoup
import urllib3
from docx import Document
import subprocess
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import shutil
# from docx2pdf import convert
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
# Obtain bot token from environment variables
BOT_TOKEN = '6610552950:AAGKhOUFb39zz1-d9W6N-jvPSt8qxMLMiUU'
bot = telebot.TeleBot(BOT_TOKEN)

user_states = {}

# States Enum


class UserStates:
    AWAITING_URL = 1


def scrapeFunction(url, bot, chat_id):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 6.3; Win 64 ; x64) Apple WeKit /537.36(KHTML , like Gecko) Chrome/80.0.3987.162 Safari/537.36'}
    foldername = url.split(sep='/')[3]

    questionsList = []  # To store the questions
    optionList = []  # To store the options list
    answerList = []  # To store the answers
    sectionList = []  # To store all the sections

    webpage = requests.get(url, headers=headers, verify=False).text
    soup = BeautifulSoup(webpage, 'lxml')
    soup.prettify()
    allsection = soup.find_all(class_='need-ul-filter')[0].find_all('li')
    for eachsection in allsection:
        sectionList.append(url[:url.find('questions-and-answers')]+str(
            eachsection.text.strip()).lower().replace("'", '').replace('.', '').replace(' ', '-')+'/')

    def fetchData(soup):
        questionStatement = soup.find_all(
            class_='bix-td-qtxt table-responsive w-100')
        optionContainer = soup.find_all(class_='bix-tbl-options')
        # Extract Question Statement
        for i in questionStatement:
            questionsList.append(i.text.strip())
        # Extract the options
        for container in optionContainer:
            options = []
            for j in container.find_all(class_='bix-td-option-val d-flex flex-row align-items-center'):
                fraction_table = j.find('table', class_='ga-tbl-answer')

                if fraction_table:  # Check if fraction structure is present
                    divident_row = fraction_table.find(
                        'tr', class_='ga-tr-divident')
                    numerator_td = divident_row.find(
                        'td', class_='ga-td-divident') if divident_row else None
                    numerator = numerator_td.text.strip() if numerator_td else ""

                    divisor_row = fraction_table.find(
                        'tr', class_='ga-tr-divisor')
                    denominator_td = divisor_row.find(
                        'td', class_='ga-td-divisor') if divisor_row else None
                    denominator = denominator_td.text.strip() if denominator_td else ""

                    main_number_td = fraction_table.find(
                        'td', class_='ga-td-line')
                    main_number = main_number_td.text.strip() if main_number_td else ""

                    option_text = f"{main_number} {numerator}/{denominator} times" if all(
                        [main_number, numerator, denominator]) else j.text.strip()
                else:
                    option_text = j.text.strip()

                options.append(option_text)
            optionList.append(options)

    # Extract the answers
        correct_answers_inputs = soup.find_all('input', class_='jq-hdnakq')
        for input_field in correct_answers_inputs:
            value = input_field['value']
            if value:
                answerList.append(f"{value}")

    def mainFunction(URL):
        # For the first page we will scrape like this
        webpage = requests.get(URL, headers=headers, verify=False).text
        print(f"Fetching: {URL}")
        bot.send_message(
            chat_id, f"*Fetching:* \n`{URL}`", parse_mode="Markdown")
        soup = BeautifulSoup(webpage, 'lxml')
        soup.prettify()
        fetchData(soup)
        try:
            second_page_url = str(soup.find_all(class_='page-item')[3])
            slicing_index = int(second_page_url.find('href'))
            urlprefix = second_page_url[slicing_index:].split('"')[
                1].split('/')[-1][:-2]
            pages = int(
                soup.find(class_='breadcrumb-item active').text.strip().split()[-1])

            # Then from page 2 we will do like this
            for i in range(2, pages+1):
                if i <= 9:
                    # please check the url sequence and change it like that format
                    new_url = URL+f'{urlprefix}0{i}'
                else:
                    # please check the url sequence and change it like that format
                    new_url = URL+f'{urlprefix}{i}'
                webpage = requests.get(
                    new_url, headers=headers, verify=False).text
                # print(f"Fetching: {new_url}") # Uncomment this line if you want to check all the urls
                soup = BeautifulSoup(webpage, 'lxml')
                soup.prettify()
                fetchData(soup)
        except:
            pass

    def makeFile(URL, bot, chat_id):
        # Create a new Document
        doc = Document()
        filename = URL.split(sep='/')[4]
        # Set default font to "Times New Roman"
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Iterate through questions and options and add them to the document
        for index, (question, options) in enumerate(zip(questionsList, optionList), 1):
            # Add the question number and question in bold
            para = doc.add_paragraph()
            run = para.add_run(f"{index}. {question}")
            run.bold = True

            # Align the question to the center
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # For the options
            option_labels = ['a', 'b', 'c', 'd']
            for label, option in zip(option_labels, options):
                doc.add_paragraph(f"{label}. {option}")

        # Add answers at the end of the document
        doc.add_paragraph("\nAnswers:")
        para = doc.add_paragraph()
        for i, answer in enumerate(answerList):
            para.add_run(f"{i+1}. (")
            run = para.add_run(answer)
            run.bold = True
            para.add_run(") ")

        # Make the folder if not exist
        if not os.path.exists(foldername):
            os.makedirs(foldername)
        # Save the document
        doc.save(f'{foldername}/{filename}.docx')
        # This will convert your DOCX file to PDF using LibreOffice
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf",
                       f"{foldername}/{filename}.docx", "--outdir", foldername])

        doc_file = f'{foldername}/{filename}.docx'
        # convert(doc_file);
        with open(doc_file, 'rb') as docx_file:
            bot.send_document(chat_id, docx_file,
                              caption=f"doc file of {filename} has been processed.")

        pdf_filename = f"{foldername}/{filename}.pdf"
        with open(pdf_filename, 'rb') as pdf_file:
            bot.send_document(chat_id, pdf_file,
                              caption=f"PDF of {filename} has been processed.")

        os.remove(pdf_filename)  # delete the PDF
        os.remove(doc_file)  # delete the DOCX

        # Now clear all the question, answer and option List
        questionsList.clear()
        optionList.clear()
        answerList.clear()

    for eachsection in sectionList:
        mainFunction(eachsection)
        makeFile(eachsection, bot, chat_id)


@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot_description = ("This bot can scrape the Indiabix website's sections such as Verbal Ability and Aptitude "
                       "and send back the questions in PDF and Word formats with answers. Now, you can print "
                       "and practice more efficiently.")

    welcome_text = f"Hey *{message.chat.first_name}*, welcome to *Indiabix Scrapper Bot*.\n\n{bot_description}"
    # Create inline keyboard buttons
    markup = types.InlineKeyboardMarkup(row_width=2)
    help_button = types.InlineKeyboardButton("Help", callback_data="help")
    sendurl_button = types.InlineKeyboardButton(
        "Send URL", callback_data="sendurl")
    demo_button = types.InlineKeyboardButton(
        "Demo", callback_data="demo")  # Added demo button
    markup.add(help_button, sendurl_button, demo_button)
    bot.reply_to(message, welcome_text,
                 reply_markup=markup, parse_mode="Markdown")
    print(f"Sent welcome to {message.chat.first_name}")


@bot.callback_query_handler(func=lambda call: True)
def callback_query(call):
    if call.data == "help":
        help_menu(call.message)
    elif call.data == "sendurl":
        request_url(call.message)
    elif call.data == "demo":
        send_demo(call.message)

@bot.message_handler(commands=['demo'])
def send_demo(message):
    demo_text = (
        "1. Click on the *Send URL* button\n"
        "2. Provide the URL of the Indiabix website section you're interested in.\n\n"
        "*Proper Format* ✅:\n"
        "`https://www.indiabix.com/aptitude/questions-and-answers/`\n"
        "\n`https://www.indiabix.com/data-interpretation/questions-and-answers/`\n"
        "\n`https://www.indiabix.com/verbal-ability/questions-and-answers/`\n"
        "\n`https://www.indiabix.com/logical-reasoning/questions-and-answers/`\n"
        "\n*Wrong Format* ❌:\n"
        "`https://www.indiabix.com/logical-reasoning/number-series/`\n"
        "\n`https://www.indiabix.com/placement-papers/companies/`\n"
        "\n`https://www.indiabix.com/online-test/aptitude-test/`\n"
        "\n`https://www.indiabix.com/puzzles/sudoku/`"
    )
    markup = types.InlineKeyboardMarkup(row_width=2)
    sendurl_button = types.InlineKeyboardButton(
        "Send URL", callback_data="sendurl")
    markup.add(sendurl_button)
    bot.send_message(chat_id=message.chat.id, text=demo_text, parse_mode='Markdown',reply_markup=markup)


@bot.message_handler(commands=['sendurl'])
def request_url(message):
    bot.reply_to(message, "Please send the URL to process.")
    # Set the user state to await the URL
    user_states[message.chat.id] = UserStates.AWAITING_URL


@bot.message_handler(func=lambda msg: user_states.get(msg.chat.id) == UserStates.AWAITING_URL and 'http' in msg.text)
def process_url(message):
    del user_states[message.chat.id]  # Clear the user's state after processing
    url = message.text
    bot.send_message(
        message.chat.id, 'Please wait your files are being processed...')
    # now calling scrapeFunction with bot and chat_id as arguments
    scrapeFunction(url, bot, message.chat.id)


@bot.message_handler(commands=['help'])
def help_menu(message):
    help_text = ("/start - Show the welcome message\n"
                 "/demo - How the bot works\n"
                 "/sendurl - Send the URL to Scrape\n"
                 "/help - Show this help menu")
    bot.send_message(message.chat.id, help_text)


@bot.message_handler(func=lambda msg: True)
def default_reply(message):
    bot.reply_to(
        message, f"Sorry {message.chat.first_name}, the bot can only work with Indiabix URL.")
    print(f"Replied to message from {message.chat.first_name}")


print("Bot started and awaiting messages...")
bot.infinity_polling()
